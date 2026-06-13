import io
import logging
import os
from typing import Optional, Tuple

import pdfplumber
from PyPDF2 import PdfReader
from PyPDF2.errors import EmptyFileError, PdfReadError
from docx import Document

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


class DocumentReadError(Exception):
    pass


def validate_file_size(file_path: str) -> None:
    if not os.path.exists(file_path):
        raise DocumentReadError(f"File not found: {file_path}")
    size = os.path.getsize(file_path)
    if size > MAX_FILE_SIZE_BYTES:
        raise DocumentReadError(
            f"File size {size / (1024 * 1024):.1f} MB exceeds maximum allowed size of {MAX_FILE_SIZE_MB} MB"
        )


def detect_file_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    if ext == ".txt":
        return "txt"
    raise DocumentReadError(f"Unsupported file type: {ext}")


def read_pdf(file_path: str) -> str:
    text_parts = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as exc:
        logger.warning("pdfplumber failed for %s, falling back to PyPDF2: %s", file_path, exc)
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except (EmptyFileError, PdfReadError) as exc2:
            raise DocumentReadError(f"Failed to read corrupted PDF: {exc2}")

    return "\n".join(text_parts)


def read_txt(file_path: str) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding, errors="strict") as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def read_document(file_path: str) -> str:
    validate_file_size(file_path)
    file_type = detect_file_type(file_path)
    logger.info("Reading %s file: %s", file_type, file_path)

    try:
        if file_type == "pdf":
            return read_pdf(file_path)
        if file_type == "docx":
            return read_docx(file_path)
        if file_type == "txt":
            return read_txt(file_path)
    except Exception as exc:
        logger.error("Error reading %s: %s", file_path, exc)
        raise DocumentReadError(f"Failed to read document: {exc}")

    raise DocumentReadError(f"Unsupported file type: {file_type}")


def extract_text_from_stream(uploaded_file: object, file_type: str) -> str:
    try:
        if file_type == "txt":
            content = uploaded_file.read()
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return content.decode("utf-8", errors="replace")
        if file_type == "pdf":
            with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n".join(text_parts)
        if file_type == "docx":
            doc = Document(io.BytesIO(uploaded_file.read()))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        raise DocumentReadError(f"Unsupported file type: {file_type}")
    except Exception as exc:
        logger.error("Error extracting text from stream: %s", exc)
        raise DocumentReadError(f"Failed to extract text: {exc}")
