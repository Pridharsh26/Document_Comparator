import os
from pypdf import PdfReader
from docx import Document


def load_document(file):

    extension = os.path.splitext(file.name)[1]

    if extension == ".pdf":

        pdf = PdfReader(file)

        text = ""

        for page in pdf.pages:
            text += page.extract_text() + "\n"

        return text

    elif extension == ".docx":

        doc = Document(file)

        text = "\n".join(
            para.text for para in doc.paragraphs
        )

        return text

    elif extension == ".txt":

        return file.read().decode("utf-8")

    else:
        raise ValueError("Unsupported file format")

# from pypdf import PdfReader
# from docx import Document
# import os


# def load_document(file_path):

#     extension = os.path.splitext(file_path)[1].lower()

#     if extension == ".pdf":

#         reader = PdfReader(file_path)

#         text = ""

#         for page in reader.pages:
#             text += page.extract_text() + "\n"

#         return text

#     elif extension == ".docx":

#         doc = Document(file_path)

#         return "\n".join(
#             para.text for para in doc.paragraphs
#         )

#     elif extension == ".txt":

#         with open(file_path, "r", encoding="utf-8") as f:
#             return f.read()

#     else:
#         raise ValueError("Unsupported file format")